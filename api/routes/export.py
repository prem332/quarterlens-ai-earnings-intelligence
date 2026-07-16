import json
import io
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from azure_clients.blob_client import BlobClient

router = APIRouter(prefix="/api/export", tags=["export"])

CONTAINER = "raw-documents"
REPORTS_PREFIX = "reports"


def _blob() -> BlobClient:
    return BlobClient()


def _load(run_id: str) -> dict:
    blob = _blob()
    path = f"{REPORTS_PREFIX}/{run_id}.json"
    if not blob.blob_exists(CONTAINER, path):
        raise HTTPException(status_code=404, detail="Run not found")
    doc = json.loads(blob.download_blob(CONTAINER, path))
    if doc.get("status") != "completed":
        raise HTTPException(status_code=409, detail="Analysis not yet complete")
    return doc


def _text(doc: dict) -> str:
    lines = [
        "QuarterLens AI — Earnings Analysis",
        f"Company: {doc['company']}  |  Quarter: {doc['quarter']}",
        f"Run ID: {doc['run_id']}",
        "",
        "=" * 60,
        "",
        doc.get("report") or "(No report generated)",
        "",
        "=" * 60,
        "NUMERIC VALIDATIONS",
        "",
    ]
    for v in doc.get("numeric_validations", []):
        status = "✓" if v.get("verified") else "✗"
        lines.append(f"{status}  {v.get('claim')}  —  filed: {v.get('filed_value')}  stated: {v.get('stated_value')}")
    return "\n".join(lines)


@router.post("/{run_id}/pdf")
async def export_pdf(run_id: str):
    doc = _load(run_id)
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import letter

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        y = height - 50
        for line in _text(doc).split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(40, y, line[:110])
            y -= 14
        c.save()
        buf.seek(0)
    except ImportError:
        content = _text(doc).encode()
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="quarterlens_{run_id}.txt"'},
        )
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="quarterlens_{run_id}.pdf"'},
    )


@router.post("/{run_id}/docx")
async def export_docx(run_id: str):
    doc = _load(run_id)
    try:
        from docx import Document

        document = Document()
        document.add_heading("QuarterLens AI — Earnings Analysis", 0)
        document.add_paragraph(f"Company: {doc['company']}   Quarter: {doc['quarter']}")
        document.add_paragraph(f"Run ID: {doc['run_id']}")
        document.add_heading("Report", level=1)
        document.add_paragraph(doc.get("report") or "(No report generated)")
        if doc.get("numeric_validations"):
            document.add_heading("Numeric Validations", level=1)
            for v in doc["numeric_validations"]:
                status = "✓" if v.get("verified") else "✗"
                document.add_paragraph(f"{status}  {v.get('claim')}  —  filed: {v.get('filed_value')}  stated: {v.get('stated_value')}")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
    except ImportError:
        content = _text(doc).encode()
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="quarterlens_{run_id}.txt"'},
        )
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="quarterlens_{run_id}.docx"'},
    )